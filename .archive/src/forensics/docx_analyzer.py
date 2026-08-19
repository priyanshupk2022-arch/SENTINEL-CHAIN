"""DOCX Forensics & Hidden Layer Parsing Engine with Header, Footer, and Table Inspection."""
import io
import docx
from typing import Tuple, List, Dict, Any
from app.models.schemas import ScanFinding
from app.forensics.unicode_sanitizer import UnicodeSanitizer

class DOCXAnalyzer:
    """
    Advanced DOCX Forensic Inspector.
    Dissects Word OpenXML structures to detect:
    1. Vanished text attributes (w:vanish / font.hidden).
    2. Invisible font colors (#FFFFFF, #FEFEFE, #FDFDFD).
    3. Sub-pixel / Micro-fonts (< 2.0 pt).
    4. Steganography hidden in Headers, Footers, and Table Cells.
    5. Prompt injections in Core Document Properties (Author, Comments, Title).
    """
    @staticmethod
    def inspect(file_bytes: bytes) -> Tuple[str, List[ScanFinding], Dict[str, Any]]:
        findings: List[ScanFinding] = []
        extracted_text: List[str] = []
        metadata: Dict[str, Any] = {}
        
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            
            # 1. Core Properties Inspection
            core_props = doc.core_properties
            metadata = {
                "author": core_props.author or "",
                "title": core_props.title or "",
                "comments": core_props.comments or "",
                "keywords": core_props.keywords or "",
                "subject": core_props.subject or "",
                "category": core_props.category or ""
            }
            
            for key, val in metadata.items():
                if val:
                    val_lower = val.lower()
                    if any(bad in val_lower for bad in [
                        "ignore previous", "system override", "<|im_start|>",
                        "admin override", "developer mode", "disregard instructions"
                    ]):
                        findings.append(ScanFinding(
                            category="metadata_injection",
                            severity="CRITICAL",
                            description=f"Prompt injection payload found in DOCX metadata field '{key}'.",
                            location=f"DOCX Metadata: {key}",
                            original_snippet=val
                        ))

            def _process_paragraph(p, context="Paragraph"):
                for run in p.runs:
                    text = run.text.strip()
                    if not text:
                        continue
                    
                    # A. Hidden text check in Word run properties (w:vanish)
                    if run.font.hidden:
                        findings.append(ScanFinding(
                            category="hidden_text",
                            severity="CRITICAL",
                            description=f"Hidden text property (w:vanish) detected in DOCX {context.lower()}.",
                            location=context,
                            original_snippet=text
                        ))
                        
                    # B. Font color check (white text in Word: FFFFFF or white)
                    font_color = run.font.color
                    if font_color and font_color.rgb:
                        rgb_hex = str(font_color.rgb).upper()
                        if rgb_hex in ("FFFFFF", "FEFEFE", "FDFDFD"):
                            findings.append(ScanFinding(
                                category="white_text",
                                severity="CRITICAL",
                                description=f"White font color (#{rgb_hex}) detected in DOCX {context.lower()}.",
                                location=context,
                                original_snippet=text
                            ))

                    # C. Micro-font check in DOCX (< 2 pt)
                    if run.font.size and run.font.size.pt < 2.0:
                        findings.append(ScanFinding(
                            category="micro_font",
                            severity="HIGH",
                            description=f"Microscopic font size ({run.font.size.pt:.1f}pt) detected in DOCX {context.lower()}.",
                            location=context,
                            original_snippet=text
                        ))
                    
                    extracted_text.append(text)

            # 2. Inspect Body Paragraphs (capped to prevent CPU exhaustion)
            MAX_DOCX_PARAGRAPHS = 5000
            if len(doc.paragraphs) > MAX_DOCX_PARAGRAPHS:
                findings.append(ScanFinding(
                    category="resource_exhaustion",
                    severity="HIGH",
                    description=f"DOCX paragraph count ({len(doc.paragraphs)}) exceeds maximum security limit ({MAX_DOCX_PARAGRAPHS}). Processing capped.",
                    location="Document Structure",
                    original_snippet=f"Total paragraphs: {len(doc.paragraphs)}"
                ))

            for i, p in enumerate(doc.paragraphs[:MAX_DOCX_PARAGRAPHS]):
                _process_paragraph(p, context=f"Paragraph {i+1}")


            # 3. Inspect Table Cells
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p in cell.paragraphs:
                            _process_paragraph(p, context=f"Table {t_idx+1} [R{r_idx+1}:C{c_idx+1}]")

            # 4. Inspect Headers & Footers
            for s_idx, section in enumerate(doc.sections):
                for p in section.header.paragraphs:
                    _process_paragraph(p, context=f"Section {s_idx+1} Header")
                for p in section.footer.paragraphs:
                    _process_paragraph(p, context=f"Section {s_idx+1} Footer")

            full_text = "\n".join(extracted_text)
            sanitized_text, u_findings = UnicodeSanitizer.sanitize(full_text)
            findings.extend(u_findings)
            
            return sanitized_text, findings, metadata
            
        except Exception as e:
            findings.append(ScanFinding(
                category="docx_parsing_error",
                severity="HIGH",
                description=f"DOCX forensic parsing encountered an anomaly: {str(e)}"
            ))
            return "", findings, metadata

# Cross-import alias
DOCXInspector = DOCXAnalyzer
